from pathlib import Path

from docx import Document

from app.ingestion.parser import ParsedDocument


class DocxParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> ParsedDocument:
        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return ParsedDocument(text=text, source_type="docx")
